# -*- coding: utf-8 -*-
"""
Generador de documentos .docx para el proyecto Zajuna.
Crea:
 - docs/Informe_Tecnico.docx
 - docs/Manual_Despliegue.docx

Requiere: python-docx (pip install python-docx)
"""
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
import os

root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
outdir = os.path.join(root, 'docs')
os.makedirs(outdir, exist_ok=True)

# Contenido del Informe técnico
informe = Document()
informe.add_heading('Informe técnico de cambios - Módulo Foro', level=1)

informe.add_heading('1) Resumen ejecutivo', level=2)
informe.add_paragraph(
    'Se modificó la lógica de notificaciones que se muestra en la vista de un tema de discusión del módulo foro ' 
    'para cumplir la regla solicitada: mostrar la fecha vigente (duedate o cutoffdate) y, si existen ambas, la mayor. ' 
    'Además se añadió distinción para que el texto indique "fecha límite" cuando corresponda.'
)

informe.add_heading('2) Archivos modificados', level=2)
informe.add_paragraph('mod/forum/classes/local/renderers/discussion.php')
informe.add_paragraph(
    'Descripción: Se actualizó la función get_notifications() para: comprobar si se alcanzó cutoffdate, elegir la fecha a mostrar ' 
    '(max entre duedate y cutoffdate) y seleccionar la cadena de idioma apropiada (fecha límite o fecha de entrega).')

informe.add_heading('3) Nuevas cadenas de idioma (recomendadas)', level=2)
informe.add_paragraph(
    'Se recomienda añadir las siguientes cadenas en mod/forum/lang/es/forum.php:\n'
    "- thisforumiscutoff = 'La fecha límite para publicar en este foro fue {$a}.'\n"
    "- thisforumhascutoffdate = 'La fecha límite para publicar en este foro es {$a}.'\n"
    'Las cadenas de due date ya existen: thisforumisdue / thisforumhasduedate.'
)

informe.add_heading('4) Comportamiento verificado', level=2)
informe.add_paragraph(
    'Pruebas manuales realizadas: al seleccionar ambas fechas se muestra la mayor; al seleccionar sólo cutoff se muestra la fecha límite; ' 
    'si cutoff está en el pasado se muestra el mensaje de corte (cutoffdatereached).' )

informe.add_heading('5) Quality gates / estado actual', level=2)
informe.add_paragraph(
    'Recomendaciones: ejecutar php -l en el archivo modificado y purgar cachés de Moodle. ' 
    'Se detectó un conflicto en un plugin local (local/forumreplyembed) que redeclara una clase renderer; revisar ese plugin si provoca errores.'
)

informe.add_heading('6) Entregables', level=2)
informe.add_paragraph('- Código modificado: mod/forum/classes/local/renderers/discussion.php')
informe.add_paragraph('- Recomendación: añadir cadenas de idioma en mod/forum/lang/es/forum.php')

informe.add_paragraph('\nDocumento generado automáticamente.')

informe_path = os.path.join(outdir, 'Informe_Tecnico.docx')
informe.save(informe_path)
print('Generado:', informe_path)

# Contenido del Manual de despliegue
manual = Document()
manual.add_heading('Manual de despliegue - Cambios en módulo Foro', level=1)

manual.add_heading('Preparación / backups', level=2)
manual.add_paragraph('Antes de desplegar, realizar backup de archivos y de la base de datos. ' 
                     'Ejecutar mysqldump o la herramienta equivalente para la BD y comprimir el directorio del sitio.')

manual.add_heading('Archivos a desplegar', level=2)
manual.add_paragraph(' - mod/forum/classes/local/renderers/discussion.php (modificado)')
manual.add_paragraph(' - mod/forum/lang/es/forum.php (opcional, añadir cadenas en español para cutoff)')

manual.add_heading('Pasos de despliegue (PowerShell)', level=2)
manual.add_paragraph('1) Hacer backup de archivos y BD')
manual.add_paragraph("2) Actualizar los ficheros en el servidor (por git, scp o deploy automático).")
manual.add_paragraph('3) Purga de cachés de Moodle:')
manual.add_paragraph('```powershell\ncd C:\\wamp64\\www\\zajuna\nphp .\\admin\\cli\\purge_caches.php\n```')

manual.add_heading('Comprobaciones post-despliegue', level=2)
manual.add_paragraph(' - Verificar en un foro de prueba: solo due date, solo cutoff, ambas (cutoff > due), cutoff en pasado.')
manual.add_paragraph(' - Revisar logs de PHP/Apache por errores.')

manual.add_heading('Rollback', level=2)
manual.add_paragraph('Si hace falta revertir: usar git revert o restaurar los backups de archivos y BD.')

manual.add_heading('Comandos útiles', level=2)
manual.add_paragraph('Comprobar sintaxis del PHP modificado:')
manual.add_paragraph('```powershell\nphp -l .\\mod\\forum\\classes\\local\\renderers\\discussion.php\n```')
manual.add_paragraph('Purgar cachés:')
manual.add_paragraph('```powershell\nphp .\\admin\\cli\\purge_caches.php\n```')

manual.add_heading('Riesgos conocidos', level=2)
manual.add_paragraph(' - Plugin local `local/forumreplyembed` detectado en el workspace puede redeclarar la clase renderer. Revisar y corregir si provoca errores.')

manual.add_paragraph('\nDocumento generado automáticamente.')

manual_path = os.path.join(outdir, 'Manual_Despliegue.docx')
manual.save(manual_path)
print('Generado:', manual_path)
